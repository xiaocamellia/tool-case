"""
试验记录助手 - Experiment Recording Assistant
基于 PyQt5，用于外场/暗室试验的结构化记录与报告导出
"""

import os
import json
import datetime
import hashlib
from PyQt5.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QGridLayout, QGroupBox,
    QPushButton, QLabel, QFileDialog, QComboBox, QDoubleSpinBox,
    QSpinBox, QMessageBox, QTextEdit, QTabWidget, QSplitter,
    QFrame, QListWidget, QListWidgetItem, QAbstractItemView,
    QLineEdit, QDialog, QFormLayout, QDialogButtonBox, QDateEdit,
    QScrollArea, QMenu, QAction, QInputDialog, QCheckBox
)
from PyQt5.QtCore import Qt, QDate, QTimer
from PyQt5.QtGui import QFont, QColor, QPixmap
from datetime import date

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ModuleNotFoundError:
    Document = None
    Inches = Pt = Cm = Emu = lambda value: value
    RGBColor = lambda red, green, blue: (red, green, blue)
    WD_ALIGN_PARAGRAPH = WD_LINE_SPACING = WD_TABLE_ALIGNMENT = WD_ORIENT = None
    qn = OxmlElement = None

from .ui_theme import (
    set_button_icon, style_primary_button, style_warning_button, style_calc_button
)
from .logger import get_logger

logger = get_logger('exp_recorder')

DATA_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), '试验数据')
os.makedirs(DATA_DIR, exist_ok=True)


class ExpRecorderPanel(QWidget):
    """试验记录助手主面板"""

    def __init__(self, parent=None):
        super().__init__(parent)
        self.projects = {}      # project_id -> project data
        self.current_project_id = None
        self.current_task_id = None
        self.current_record_id = None
        self._load_data()
        self._init_ui()
        logger.info('试验记录助手已创建')

    def _init_ui(self):
        layout = QVBoxLayout(self)
        layout.setContentsMargins(6, 6, 6, 6)

        # ===== 顶部工具栏 =====
        top_layout = QHBoxLayout()
        lbl_title = QLabel('📋 试验记录助手')
        lbl_title.setStyleSheet('font-size: 18px; font-weight: bold; color: #1e293b;')
        top_layout.addWidget(lbl_title)

        self.btn_new_project = QPushButton('📁 新建项目')
        self.btn_new_project.clicked.connect(self._new_project)
        style_primary_button(self.btn_new_project)
        top_layout.addWidget(self.btn_new_project)

        self.btn_new_task = QPushButton('📂 新建任务')
        self.btn_new_task.clicked.connect(self._new_task)
        top_layout.addWidget(self.btn_new_task)

        self.btn_new_record = QPushButton('📝 新建记录')
        self.btn_new_record.clicked.connect(self._new_record)
        style_primary_button(self.btn_new_record)
        top_layout.addWidget(self.btn_new_record)

        top_layout.addStretch()
        self.lbl_save_status = QLabel('')
        self.lbl_save_status.setStyleSheet('font-size: 10px; color: #94a3b8;')
        top_layout.addWidget(self.lbl_save_status)
        self.btn_save = QPushButton('💾 保存')
        self.btn_save.clicked.connect(self._save_current_record)
        style_primary_button(self.btn_save)
        top_layout.addWidget(self.btn_save)

        # 导出菜单
        self.btn_export_word = QPushButton('📤 导出 ▼')
        self.btn_export_word.clicked.connect(self._show_export_menu)
        style_calc_button(self.btn_export_word)
        self.export_menu = QMenu(self)
        act_export_record = QAction('📝 导出当前记录', self)
        act_export_record.triggered.connect(lambda: self._export_to_word('record'))
        self.export_menu.addAction(act_export_record)
        act_export_task = QAction('📂 导出当前任务（含所有记录）', self)
        act_export_task.triggered.connect(lambda: self._export_to_word('task'))
        self.export_menu.addAction(act_export_task)
        act_export_project = QAction('📁 导出当前项目（含所有任务和记录）', self)
        act_export_project.triggered.connect(lambda: self._export_to_word('project'))
        self.export_menu.addAction(act_export_project)
        top_layout.addWidget(self.btn_export_word)

        # 树右键菜单
        self.tree_context_menu = QMenu(self)
        act_export_rec = QAction('📝 导出此记录', self)
        act_export_rec.triggered.connect(lambda: self._export_tree_selected('record'))
        self.tree_context_menu.addAction(act_export_rec)
        act_export_tsk = QAction('📂 导出此任务', self)
        act_export_tsk.triggered.connect(lambda: self._export_tree_selected('task'))
        self.tree_context_menu.addAction(act_export_tsk)
        act_export_prj = QAction('📁 导出此项目', self)
        act_export_prj.triggered.connect(lambda: self._export_tree_selected('project'))
        self.tree_context_menu.addAction(act_export_prj)
        self.tree_context_menu.addSeparator()
        act_delete_tree = QAction('🗑️ 删除', self)
        act_delete_tree.triggered.connect(self._delete_selected)
        self.tree_context_menu.addAction(act_delete_tree)
        layout.addLayout(top_layout)

        # ===== 主分割器 =====
        splitter = QSplitter(Qt.Horizontal)

        # ---- 左侧项目树 ----
        left_widget = QWidget()
        left_layout = QVBoxLayout(left_widget)
        left_layout.setContentsMargins(0, 0, 0, 0)

        left_layout.addWidget(QLabel('📁 项目结构'))
        self.tree_projects = QListWidget()
        self.tree_projects.setMinimumWidth(200)
        self.tree_projects.currentItemChanged.connect(self._on_tree_selection)
        self.tree_projects.setContextMenuPolicy(Qt.CustomContextMenu)
        self.tree_projects.customContextMenuRequested.connect(self._show_tree_context_menu)
        left_layout.addWidget(self.tree_projects, 1)

        self.btn_delete_item = QPushButton('🗑️ 删除选中')
        self.btn_delete_item.clicked.connect(self._delete_selected)
        style_warning_button(self.btn_delete_item)
        left_layout.addWidget(self.btn_delete_item)

        splitter.addWidget(left_widget)

        # ---- 右侧编辑区 ----
        right_widget = QWidget()
        right_layout = QVBoxLayout(right_widget)
        right_layout.setContentsMargins(0, 0, 0, 0)

        # 页面切换
        self.stacked = QTabWidget()

        # 欢迎页面 - 功能特色介绍
        self.page_empty = QWidget()
        empty_layout = QVBoxLayout(self.page_empty)
        empty_layout.setContentsMargins(30, 30, 30, 30)

        # 标题区
        title_lbl = QLabel('📋 试验记录助手')
        title_lbl.setStyleSheet('font-size: 24px; font-weight: bold; color: #1e293b;')
        title_lbl.setAlignment(Qt.AlignCenter)
        empty_layout.addWidget(title_lbl)

        sub_lbl = QLabel('告别Word排版 · 一键生成规范试验报告')
        sub_lbl.setStyleSheet('font-size: 13px; color: #64748b; margin-bottom: 20px;')
        sub_lbl.setAlignment(Qt.AlignCenter)
        empty_layout.addWidget(sub_lbl)

        # 核心功能卡片
        features = [
            ('📁 三级项目管理', '项目 → 任务 → 记录，结构化组织试验数据'),
            ('📝 模板化记录', '暗室天线测试 / 外场雷达试验 / 系统联调 / 通用，填空式录入'),
            ('🖼️ 截图管理', '支持多种图片格式，自定义图注说明，自动排序编号'),
            ('📤 一键导出Word', '军工归档格式：标准封面 + 页眉页码 + 目录 + 图题注 + 尾注'),
            ('⚙️ 导出设置', '封面 / 页码 / 自动目录 / 图片质量 / 水印 可自由配置'),
            ('📂 数据持久化', '本地JSON存储，自动保存，重启不丢失'),
            ('🔍 快速导航', '左侧树形结构，右键菜单支持导出/删除'),
            ('🎯 工科规范', '宋体正文/黑体标题/1.5倍行距/首行缩进/纯黑灰配色'),
        ]

        # 用表格布局展示功能卡片
        from PyQt5.QtWidgets import QGridLayout
        grid = QGridLayout()
        grid.setSpacing(12)
        for i, (icon_title, desc) in enumerate(features):
            row, col = i // 2, i % 2
            card = QFrame()
            card.setStyleSheet("""
                QFrame {
                    background: #ffffff;
                    border: 1px solid #e2e8f0;
                    border-radius: 10px;
                    padding: 14px;
                }
                QFrame:hover {
                    border-color: #94a3b8;
                    background: #f8fafc;
                }
            """)
            card_layout = QVBoxLayout(card)
            card_layout.setContentsMargins(12, 10, 12, 10)
            card_layout.setSpacing(4)
            icon_lbl = QLabel(icon_title)
            icon_lbl.setStyleSheet('font-size: 14px; font-weight: bold; color: #1e293b;')
            card_layout.addWidget(icon_lbl)
            desc_lbl = QLabel(desc)
            desc_lbl.setStyleSheet('font-size: 11px; color: #64748b;')
            desc_lbl.setWordWrap(True)
            card_layout.addWidget(desc_lbl)
            grid.addWidget(card, row, col)

        empty_layout.addLayout(grid)

        # 底部操作提示
        empty_layout.addSpacing(20)
        tip_lbl = QLabel('💡 点击左侧「📁 项目结构」→ 右键或顶部按钮新建项目/任务/记录')
        tip_lbl.setStyleSheet('font-size: 12px; color: #94a3b8; padding: 8px; background: #f1f5f9; border-radius: 6px;')
        tip_lbl.setAlignment(Qt.AlignCenter)
        empty_layout.addWidget(tip_lbl)

        empty_layout.addStretch()
        self.stacked.addTab(self.page_empty, '欢迎')

        # 记录编辑页
        self.page_edit = QScrollArea()
        self.page_edit.setWidgetResizable(True)
        self.edit_container = QWidget()
        self.edit_layout = QVBoxLayout(self.edit_container)
        self._build_edit_form()
        self.page_edit.setWidget(self.edit_container)
        self.stacked.addTab(self.page_edit, '编辑记录')

        right_layout.addWidget(self.stacked, 1)
        splitter.addWidget(right_widget)
        splitter.setStretchFactor(0, 1)
        splitter.setStretchFactor(1, 3)

        layout.addWidget(splitter, 1)
        self._refresh_tree()

    def _build_edit_form(self):
        """构建记录编辑表单"""
        self.edit_form_widgets = {}

        # 标题
        self.edit_title = QLineEdit()
        self.edit_title.setPlaceholderText('请输入试验记录标题')
        self.edit_title.setStyleSheet('font-size: 16px; font-weight: bold; padding: 6px;')
        self.edit_layout.addWidget(self.edit_title)

        # 基本信息
        info_group = QGroupBox('基本信息')
        info_layout = QGridLayout(info_group)

        info_layout.addWidget(QLabel('试验日期:'), 0, 0)
        self.edit_date = QDateEdit()
        self.edit_date.setDate(QDate.currentDate())
        self.edit_date.setCalendarPopup(True)
        info_layout.addWidget(self.edit_date, 0, 1)

        info_layout.addWidget(QLabel('试验人员:'), 0, 2)
        self.edit_personnel = QLineEdit()
        self.edit_personnel.setPlaceholderText('姓名1, 姓名2')
        info_layout.addWidget(self.edit_personnel, 0, 3)

        info_layout.addWidget(QLabel('试验地点:'), 1, 0)
        self.edit_location = QLineEdit()
        self.edit_location.setPlaceholderText('外场/暗室/实验室')
        info_layout.addWidget(self.edit_location, 1, 1)

        info_layout.addWidget(QLabel('模板选择:'), 1, 2)
        self.combo_template = QComboBox()
        self.combo_template.addItems(['暗室天线测试', '外场雷达试验', '系统联调测试', '通用试验记录'])
        info_layout.addWidget(self.combo_template, 1, 3)

        info_layout.addWidget(QLabel('环境条件:'), 2, 0)
        self.edit_conditions = QLineEdit()
        self.edit_conditions.setPlaceholderText('天气/温度/湿度等')
        info_layout.addWidget(self.edit_conditions, 2, 1, 1, 3)

        self.edit_layout.addWidget(info_group)

        # 试验内容 - 多行文本
        content_group = QGroupBox('试验内容')
        content_layout = QVBoxLayout(content_group)
        self.edit_content = QTextEdit()
        self.edit_content.setPlaceholderText('描述试验目的、测试项目、配置信息等...')
        self.edit_content.setMinimumHeight(100)
        content_layout.addWidget(self.edit_content)
        self.edit_layout.addWidget(content_group)

        # 试验步骤（循环添加）
        self.steps_group = QGroupBox('试验步骤')
        self.steps_layout = QVBoxLayout(self.steps_group)
        self.step_widgets = []
        self._add_step()
        btn_add_step = QPushButton('➕ 添加步骤')
        btn_add_step.clicked.connect(self._add_step)
        self.steps_layout.addWidget(btn_add_step)
        self.edit_layout.addWidget(self.steps_group)

        # 结论
        concl_group = QGroupBox('结论与问题')
        concl_layout = QVBoxLayout(concl_group)
        self.edit_conclusion = QTextEdit()
        self.edit_conclusion.setPlaceholderText('试验结论、问题汇总、后续计划...')
        self.edit_conclusion.setMinimumHeight(80)
        concl_layout.addWidget(self.edit_conclusion)
        self.edit_layout.addWidget(concl_group)

        self.edit_layout.addStretch()

    def _add_step(self):
        """添加一个试验步骤（含数据记录和截图管理）"""
        step_id = len(self.step_widgets)
        frame = QFrame()
        frame.setFrameShape(QFrame.StyledPanel)
        frame.setStyleSheet('QFrame { border: 1px solid #e2e8f0; border-radius: 4px; padding: 4px; margin: 2px; }')
        step_layout = QVBoxLayout(frame)

        # 步骤标题行
        top_row = QHBoxLayout()
        lbl = QLabel(f'步骤 {step_id + 1}:')
        lbl.setStyleSheet('font-weight: bold;')
        top_row.addWidget(lbl)

        title_edit = QLineEdit()
        title_edit.setPlaceholderText('步骤名称（可选）')
        top_row.addWidget(title_edit, 1)

        btn_del = QPushButton('\u2715')
        btn_del.setFixedSize(24, 24)
        btn_del.setStyleSheet('background: #fee2e2; color: #ef4444; border: none; border-radius: 4px;')
        btn_del.clicked.connect(lambda: self._remove_step(frame))
        top_row.addWidget(btn_del)

        step_layout.addLayout(top_row)

        # 操作描述
        desc_edit = QTextEdit()
        desc_edit.setPlaceholderText('操作描述和现象...')
        desc_edit.setMaximumHeight(60)
        step_layout.addWidget(desc_edit)

        # ---- 步骤内数据记录 ----
        data_edit = QTextEdit()
        data_edit.setPlaceholderText('此步骤的测试数据、观测结果...（可粘贴Excel表格数据）')
        data_edit.setMaximumHeight(80)
        step_layout.addWidget(data_edit)

        # ---- 步骤内截图管理 ----
        img_row = QHBoxLayout()
        btn_add_img = QPushButton('\U0001f5bc\ufe0f 添加截图')
        from .ui_theme import style_primary_button, style_warning_button
        btn_add_img.clicked.connect(lambda f=frame: self._add_step_screenshot(f))
        style_primary_button(btn_add_img)
        img_row.addWidget(btn_add_img)
        btn_remove_img = QPushButton('\u2716 移除选中')
        btn_remove_img.clicked.connect(lambda f=frame: self._remove_step_screenshot(f))
        style_warning_button(btn_remove_img)
        img_row.addWidget(btn_remove_img)
        img_row.addStretch()
        step_layout.addLayout(img_row)

        img_list = QListWidget()
        img_list.setMaximumHeight(80)
        step_layout.addWidget(img_list)

        self.steps_layout.insertWidget(self.steps_layout.count() - 1, frame)
        self.step_widgets.append({
            'frame': frame,
            'title': title_edit,
            'desc': desc_edit,
            'data': data_edit,
            'img_list': img_list,
        })
        self._renumber_steps()

    def _remove_step(self, frame):
        """删除步骤"""
        for sw in self.step_widgets[:]:
            if sw['frame'] == frame:
                self.steps_layout.removeWidget(frame)
                frame.deleteLater()
                self.step_widgets.remove(sw)
        self._renumber_steps()


    def _add_step_screenshot(self, frame):
        """向指定步骤添加截图"""
        filepath, _ = QFileDialog.getOpenFileName(
            self, '选择截图', '',
            '图片文件 (*.png *.jpg *.jpeg *.bmp *.tiff *.webp);;所有文件 (*.*)')
        if not filepath:
            return
        # 找到对应步骤的 widget
        for sw in self.step_widgets:
            if sw['frame'] == frame:
                caption, ok = QInputDialog.getText(self, '添加截图', '图注说明（可选）:')
                if not ok:
                    caption = ''
                basename = os.path.basename(filepath)
                display = f'🖼️ {basename}'
                if caption:
                    display += f' → {caption}'
                item = QListWidgetItem(display)
                item.setData(Qt.UserRole, {'path': filepath, 'caption': caption})
                item.setToolTip(f'{filepath}\n图注: {caption if caption else "（无）"}')
                sw['img_list'].addItem(item)
                break

    def _remove_step_screenshot(self, frame):
        """移除指定步骤的选中截图"""
        for sw in self.step_widgets:
            if sw['frame'] == frame:
                for item in sw['img_list'].selectedItems():
                    sw['img_list'].takeItem(sw['img_list'].row(item))
                break

    def _renumber_steps(self):
        """重新编号步骤"""
        for i, sw in enumerate(self.step_widgets):
            label = sw['frame'].layout().itemAt(0).layout().itemAt(0).widget()
            if isinstance(label, QLabel):
                label.setText(f'步骤 {i + 1}:')

    # ==================== 项目/任务/记录管理 ====================

    def _refresh_tree(self):
        """刷新项目树"""
        self.tree_projects.blockSignals(True)
        self.tree_projects.clear()

        for pid, proj in sorted(self.projects.items(), key=lambda x: x[1].get('name', '')):
            proj_item = QListWidgetItem(f'📁 {proj.get("name", "未命名")}')
            proj_item.setData(Qt.UserRole, ('project', pid))
            font = proj_item.font()
            font.setBold(True)
            proj_item.setFont(font)
            self.tree_projects.addItem(proj_item)

            for tid, task in sorted(proj.get('tasks', {}).items(), key=lambda x: x[1].get('name', '')):
                task_item = QListWidgetItem(f'  📂 {task.get("name", "未命名")}')
                task_item.setData(Qt.UserRole, ('task', pid, tid))
                self.tree_projects.addItem(task_item)

                for rid, rec in sorted(task.get('records', {}).items(), key=lambda x: x[1].get('title', '')):
                    rec_item = QListWidgetItem(f'    📝 {rec.get("title", "未命名")}')
                    rec_item.setData(Qt.UserRole, ('record', pid, tid, rid))
                    self.tree_projects.addItem(rec_item)

        self.tree_projects.blockSignals(False)

    def _on_tree_selection(self, current, previous):
        """树节点选择事件"""
        if not current:
            return
        data = current.data(Qt.UserRole)
        if not data:
            return

        if data[0] == 'record':
            _, pid, tid, rid = data
            self._load_record(pid, tid, rid)
        elif data[0] in ('project', 'task'):
            self.stacked.setCurrentIndex(0)

    def _new_project(self):
        """新建项目"""
        name, ok = QInputDialog.getText(self, '新建项目', '项目名称:')
        if not ok or not name.strip():
            return
        desc, ok = QInputDialog.getText(self, '新建项目', '项目描述（可选）:')
        if not ok:
            desc = ''

        pid = hashlib.md5((name + str(datetime.datetime.now())).encode()).hexdigest()[:8]
        self.projects[pid] = {
            'name': name.strip(),
            'description': desc,
            'tasks': {}
        }
        self._save_data()
        self._refresh_tree()
        logger.info(f'新建项目: {name.strip()}')

    def _new_task(self):
        """新建任务"""
        if not self.projects:
            QMessageBox.information(self, '提示', '请先创建项目')
            return

        proj_names = [p['name'] for p in self.projects.values()]
        proj_name, ok = QInputDialog.getItem(self, '选择项目', '选择所属项目:', proj_names, 0, False)
        if not ok:
            return

        pid = None
        for k, v in self.projects.items():
            if v['name'] == proj_name:
                pid = k
                break
        if not pid:
            return

        name, ok = QInputDialog.getText(self, '新建任务', '任务名称:')
        if not ok or not name.strip():
            return

        tid = hashlib.md5((name + str(datetime.datetime.now())).encode()).hexdigest()[:8]
        self.projects[pid]['tasks'][tid] = {
            'name': name,
            'records': {}
        }
        self._save_data()
        self._refresh_tree()
        logger.info(f'新建任务: {name}')

    def _new_record(self):
        """新建记录"""
        # 找到第一个有task的项目
        pid = None
        tid = None
        # 尝试获取当前选中的task
        current = self.tree_projects.currentItem()
        if current:
            data = current.data(Qt.UserRole)
            if data and data[0] == 'task':
                _, pid, tid = data
            elif data and data[0] == 'record':
                _, pid, tid, _ = data

        if not pid or not tid:
            QMessageBox.information(self, '提示', '请先在左侧选中一个任务（或展开项目下的任务）')
            return

        title, ok = QInputDialog.getText(self, '新建记录', '记录标题:')
        if not ok or not title.strip():
            return

        rid = hashlib.md5((title + str(datetime.datetime.now())).encode()).hexdigest()[:8]

        record = {
            'title': title,
            'date': datetime.date.today().isoformat(),
            'personnel': '',
            'location': '',
            'template': '通用试验记录',
            'conditions': '',
            'content': '',
            'steps': [],
            'conclusion': '',
            'created': datetime.datetime.now().isoformat(),
        }

        self.projects[pid]['tasks'][tid]['records'][rid] = record
        self._save_data()
        self._refresh_tree()
        self._load_record(pid, tid, rid)
        logger.info(f'新建记录: {title}')

    def _load_record(self, pid, tid, rid):
        """加载记录到编辑表单"""
        try:
            rec = self.projects[pid]['tasks'][tid]['records'][rid]
        except KeyError:
            return

        self.current_project_id = pid
        self.current_task_id = tid
        self.current_record_id = rid

        self.edit_title.setText(rec.get('title', ''))
        try:
            self.edit_date.setDate(QDate.fromString(rec.get('date', datetime.date.today().isoformat()), 'yyyy-MM-dd'))
        except:
            self.edit_date.setDate(QDate.currentDate())
        self.edit_personnel.setText(rec.get('personnel', ''))
        self.edit_location.setText(rec.get('location', ''))
        idx = self.combo_template.findText(rec.get('template', '通用试验记录'))
        if idx >= 0:
            self.combo_template.setCurrentIndex(idx)
        self.edit_conditions.setText(rec.get('conditions', ''))
        self.edit_content.setText(rec.get('content', ''))
        self.edit_conclusion.setText(rec.get('conclusion', ''))

        # 步骤（含步骤内数据记录和截图）
        for sw in self.step_widgets[:]:
            self.steps_layout.removeWidget(sw['frame'])
            sw['frame'].deleteLater()
        self.step_widgets = []

        steps = rec.get('steps', [])
        if steps:
            for s in steps:
                self._add_step()
                if self.step_widgets:
                    sw = self.step_widgets[-1]
                    sw['title'].setText(s.get('title', ''))
                    sw['desc'].setText(s.get('desc', ''))
                    sw['data'].setText(s.get('data', ''))
                    # 加载步骤内截图
                    for img_data in s.get('images', []):
                        if isinstance(img_data, dict):
                            ipath = img_data.get('path', '')
                            icaption = img_data.get('caption', '')
                        elif isinstance(img_data, str):
                            ipath = img_data
                            icaption = ''
                        else:
                            continue
                        if not os.path.exists(ipath):
                            continue
                        display = f'🖼️ {os.path.basename(ipath)}'
                        if icaption:
                            display += f' → {icaption}'
                        item = QListWidgetItem(display)
                        item.setData(Qt.UserRole, {'path': ipath, 'caption': icaption})
                        item.setToolTip(f'{ipath}\n图注: {icaption if icaption else "（无）"}')
                        sw['img_list'].addItem(item)
        else:
            self._add_step()


        self.stacked.setCurrentIndex(1)
        self.lbl_save_status.setText('已加载')

    def _save_current_record(self):
        """保存当前记录（含图注信息）"""
        if not all([self.current_project_id, self.current_task_id, self.current_record_id]):
            QMessageBox.information(self, '提示', '没有正在编辑的记录')
            return

        steps_data = []
        for sw in self.step_widgets:
            step_images = []
            for i in range(sw['img_list'].count()):
                item = sw['img_list'].item(i)
                data = item.data(Qt.UserRole)
                if data:
                    if isinstance(data, str):
                        step_images.append({'path': data, 'caption': ''})
                    else:
                        step_images.append(data)
            steps_data.append({
                'title': sw['title'].text(),
                'desc': sw['desc'].toPlainText(),
                'data': sw['data'].toPlainText(),
                'images': step_images,
            })

        try:
            rec = self.projects[self.current_project_id]['tasks'][self.current_task_id]['records'][self.current_record_id]
        except KeyError:
            return

        rec['title'] = self.edit_title.text()
        rec['date'] = self.edit_date.date().toString('yyyy-MM-dd')
        rec['personnel'] = self.edit_personnel.text()
        rec['location'] = self.edit_location.text()
        rec['template'] = self.combo_template.currentText()
        rec['conditions'] = self.edit_conditions.text()
        rec['content'] = self.edit_content.toPlainText()
        rec['steps'] = steps_data
        rec['conclusion'] = self.edit_conclusion.toPlainText()
        rec['updated'] = datetime.datetime.now().isoformat()

        self._save_data()
        self._refresh_tree()
        self.lbl_save_status.setText('✅ 已保存 ' + datetime.datetime.now().strftime('%H:%M'))
        logger.info(f'保存记录: {rec["title"]}')

    def _delete_selected(self):
        """删除选中的项目/任务/记录"""
        current = self.tree_projects.currentItem()
        if not current:
            return
        data = current.data(Qt.UserRole)
        if not data:
            return

        reply = QMessageBox.question(self, '确认删除', '确定要删除吗？\n此操作不可恢复！',
                                     QMessageBox.Yes | QMessageBox.No)
        if reply != QMessageBox.Yes:
            return

        if data[0] == 'record':
            _, pid, tid, rid = data
            del self.projects[pid]['tasks'][tid]['records'][rid]
            if self.current_record_id == rid:
                self.current_record_id = None
                self.stacked.setCurrentIndex(0)
        elif data[0] == 'task':
            _, pid, tid = data
            del self.projects[pid]['tasks'][tid]
        elif data[0] == 'project':
            pid = data[1]
            del self.projects[pid]
            if self.current_project_id == pid:
                self.current_project_id = None
                self.current_task_id = None
                self.current_record_id = None
                self.stacked.setCurrentIndex(0)

        self._save_data()
        self._refresh_tree()
        logger.info('删除成功')

    # ==================== 工科报告内置样式模板 - 方案B（全黑/灰军工规范）====================

    # ---- 全局样式常量（修改此处即可全局调整，无需改动业务代码） ----
    _REPORT_STYLES = {
        # ── 页面 ──
        'page_top_margin': Cm(2.5),
        'page_bottom_margin': Cm(2.5),
        'page_left_margin': Cm(2.8),    # 左侧装订预留
        'page_right_margin': Cm(2.5),

        # ── 封面 ──
        'cover_title_font': '黑体',
        'cover_title_size': Pt(22),     # 二号
        'cover_title_color': RGBColor(0x00, 0x00, 0x00),  # 纯黑
        'cover_title_space_before': Pt(240),  # ≈12行间距
        'cover_subtitle_font': '黑体',
        'cover_subtitle_size': Pt(15),   # 小三号
        'cover_subtitle_color': RGBColor(0x66, 0x66, 0x66),  # 深灰
        'cover_subtitle_space_before': Pt(20),
        'cover_info_font': '宋体',
        'cover_info_size': Pt(12),       # 小四
        'cover_info_color': RGBColor(0x99, 0x99, 0x99),  # 浅灰
        'cover_info_space_before': Pt(60),

        # ── 一级标题 H1（项目总章节）黑体四号纯黑 ──
        'h1_font': '黑体',
        'h1_size': Pt(16),
        'h1_color': RGBColor(0x00, 0x00, 0x00),
        'h1_bold': True,
        'h1_space_before': Pt(18),
        'h1_space_after': Pt(12),

        # ── 二级标题 H2（单条记录）黑体小四纯黑 ──
        'h2_font': '黑体',
        'h2_size': Pt(14),
        'h2_color': RGBColor(0x00, 0x00, 0x00),
        'h2_bold': True,
        'h2_space_before': Pt(12),
        'h2_space_after': Pt(6),

        # ── 三级标题 H3（内容/步骤/数据/结论）黑体12pt纯黑 ──
        'h3_font': '黑体',
        'h3_size': Pt(12),
        'h3_color': RGBColor(0x00, 0x00, 0x00),
        'h3_bold': True,
        'h3_space_before': Pt(8),
        'h3_space_after': Pt(4),

        # ── 正文 Body（宋体小四、1.5倍行距、首行缩进2字符） ──
        'body_font': '宋体',
        'body_size': Pt(11),               # 小四
        'body_color': RGBColor(0x00, 0x00, 0x00),
        'body_line_spacing': 1.5,
        'body_first_line_indent': Cm(0.74), # ≈2字符
        'body_space_before': Pt(0),
        'body_space_after': Pt(0),

        # ── 试验基础信息字段（标签加粗、内容常规、无缩进紧凑） ──
        'field_font': '宋体',
        'field_size': Pt(11),
        'field_label_color': RGBColor(0x00, 0x00, 0x00),
        'field_value_color': RGBColor(0x00, 0x00, 0x00),
        'field_space_before': Pt(2),
        'field_space_after': Pt(2),

        # ── 图片题注（小五居中 #666666） ──
        'fig_caption_font': '宋体',
        'fig_caption_size': Pt(9),          # 小五
        'fig_caption_color': RGBColor(0x66, 0x66, 0x66),

        # ── 页眉（小五、居中、中度灰） ──
        'header_font': '宋体',
        'header_size': Pt(9),
        'header_color': RGBColor(0x99, 0x99, 0x99),

        # ── 页脚页码（小五、居中、中度灰） ──
        'footer_font': '宋体',
        'footer_size': Pt(9),
        'footer_color': RGBColor(0x99, 0x99, 0x99),

        # ── 尾注（浅灰小五居中） ──
        'endnote_font': '宋体',
        'endnote_size': Pt(9),
        'endnote_color': RGBColor(0x99, 0x99, 0x99),

        # ── 水印（浅灰色） ──
        'watermark_color': RGBColor(0xcc, 0xcc, 0xcc),
    }

    _IMG_WIDTH_MAP = {'高清': Inches(5.5), '标准': Inches(5.0), '压缩': Inches(4.0)}

    # ---- 内置命名样式创建（批量建立，后续全部引用样式名，无零散硬编码） ----
    @staticmethod
    def _set_cn_font(run, font_name):
        """设置中文字体（兼容WPS+Office）"""
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:eastAsia'), font_name)
        rFonts.set(qn('w:ascii'), 'Arial')
        rFonts.set(qn('w:hAnsi'), 'Arial')

    @staticmethod
    def _create_named_styles(doc):
        """在文档中批量创建预定义命名样式"""
        S = ExpRecorderPanel._REPORT_STYLES
        styles = doc.styles

        # CoverTitle 封面主标题
        s = styles.add_style('CoverTitle', 1)
        s.font.name = S['cover_title_font']
        s.font.size = S['cover_title_size']
        s.font.bold = True
        s.font.color.rgb = S['cover_title_color']
        s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        s.paragraph_format.space_before = S['cover_title_space_before']
        ExpRecorderPanel._set_font_rpr(s.element, S['cover_title_font'])

        # CoverSubtitle 封面子标题
        s = styles.add_style('CoverSubtitle', 1)
        s.font.name = S['cover_subtitle_font']
        s.font.size = S['cover_subtitle_size']
        s.font.color.rgb = S['cover_subtitle_color']
        s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        s.paragraph_format.space_before = S['cover_subtitle_space_before']
        ExpRecorderPanel._set_font_rpr(s.element, S['cover_subtitle_font'])

        # CoverInfo 封面附属信息
        s = styles.add_style('CoverInfo', 1)
        s.font.name = S['cover_info_font']
        s.font.size = S['cover_info_size']
        s.font.color.rgb = S['cover_info_color']
        s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        s.paragraph_format.space_before = S['cover_info_space_before']
        ExpRecorderPanel._set_font_rpr(s.element, S['cover_info_font'])

        # Report_H1 一级标题
        s = styles.add_style('Report_H1', 1)
        s.font.name = S['h1_font']
        s.font.size = S['h1_size']
        s.font.bold = S['h1_bold']
        s.font.color.rgb = S['h1_color']
        s.paragraph_format.space_before = S['h1_space_before']
        s.paragraph_format.space_after = S['h1_space_after']
        ExpRecorderPanel._set_font_rpr(s.element, S['h1_font'])

        # Report_H2 二级标题
        s = styles.add_style('Report_H2', 1)
        s.font.name = S['h2_font']
        s.font.size = S['h2_size']
        s.font.bold = S['h2_bold']
        s.font.color.rgb = S['h2_color']
        s.paragraph_format.space_before = S['h2_space_before']
        s.paragraph_format.space_after = S['h2_space_after']
        ExpRecorderPanel._set_font_rpr(s.element, S['h2_font'])

        # Report_H3 三级标题
        s = styles.add_style('Report_H3', 1)
        s.font.name = S['h3_font']
        s.font.size = S['h3_size']
        s.font.bold = S['h3_bold']
        s.font.color.rgb = S['h3_color']
        s.paragraph_format.space_before = S['h3_space_before']
        s.paragraph_format.space_after = S['h3_space_after']
        ExpRecorderPanel._set_font_rpr(s.element, S['h3_font'])

        # Report_Body 正文通用
        s = styles.add_style('Report_Body', 1)
        s.font.name = S['body_font']
        s.font.size = S['body_size']
        s.font.color.rgb = S['body_color']
        s.paragraph_format.first_line_indent = S['body_first_line_indent']
        s.paragraph_format.line_spacing = S['body_line_spacing']
        s.paragraph_format.space_before = S['body_space_before']
        s.paragraph_format.space_after = S['body_space_after']
        ExpRecorderPanel._set_font_rpr(s.element, S['body_font'])

        # Report_Field_Info 试验基础信息
        s = styles.add_style('Report_Field_Info', 1)
        s.font.name = S['field_font']
        s.font.size = S['field_size']
        s.font.color.rgb = S['field_value_color']
        s.paragraph_format.space_before = S['field_space_before']
        s.paragraph_format.space_after = S['field_space_after']
        ExpRecorderPanel._set_font_rpr(s.element, S['field_font'])

        # Report_Image_Caption 图片题注
        s = styles.add_style('Report_Image_Caption', 1)
        s.font.name = S['fig_caption_font']
        s.font.size = S['fig_caption_size']
        s.font.color.rgb = S['fig_caption_color']
        s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ExpRecorderPanel._set_font_rpr(s.element, S['fig_caption_font'])

        # Report_Endnote 尾注
        s = styles.add_style('Report_Endnote', 1)
        s.font.name = S['endnote_font']
        s.font.size = S['endnote_size']
        s.font.color.rgb = S['endnote_color']
        s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ExpRecorderPanel._set_font_rpr(s.element, S['endnote_font'])

        return {
            'CoverTitle': 'CoverTitle',
            'CoverSubtitle': 'CoverSubtitle',
            'CoverInfo': 'CoverInfo',
            'Report_H1': 'Report_H1',
            'Report_H2': 'Report_H2',
            'Report_H3': 'Report_H3',
            'Report_Body': 'Report_Body',
            'Report_Field_Info': 'Report_Field_Info',
            'Report_Image_Caption': 'Report_Image_Caption',
            'Report_Endnote': 'Report_Endnote',
        }

    @staticmethod
    def _set_font_rpr(style_element, font_name):
        """为样式element设置中文字体rPr"""
        rPr = style_element.find(qn('w:rPr'))
        if rPr is None:
            rPr = OxmlElement('w:rPr')
            style_element.insert(0, rPr)
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:eastAsia'), font_name)
        rFonts.set(qn('w:ascii'), 'Arial')
        rFonts.set(qn('w:hAnsi'), 'Arial')

    # ---- 导出内容生成方法（全部调用命名样式，无零散硬编码） ----

    def _pstyle(self, doc, ns, style_name):
        """快捷获取段落样式"""
        return doc.styles[ns[style_name]]

    def _add_cover_page(self, doc, ns, title_text, include_cover):
        """添加封面（黑体二号纯黑 / 副标题深灰 / 日期浅灰）"""
        if not include_cover:
            return
        # 主标题
        p = doc.add_paragraph()
        p.style = self._pstyle(doc, ns, 'CoverTitle')
        p.add_run(title_text)

        # 副标题
        p = doc.add_paragraph()
        p.style = self._pstyle(doc, ns, 'CoverSubtitle')
        p.add_run('试验记录报告')

        # 附属信息
        now_str = datetime.datetime.now().strftime('%Y年%m月%d日')
        p = doc.add_paragraph()
        p.style = self._pstyle(doc, ns, 'CoverInfo')
        p.add_run(f'导出日期：{now_str}')

        doc.add_page_break()

    def _add_h1(self, doc, ns, text):
        """添加一级标题"""
        p = doc.add_paragraph()
        p.style = self._pstyle(doc, ns, 'Report_H1')
        p.add_run(text)

    def _add_h2(self, doc, ns, text):
        """添加二级标题"""
        p = doc.add_paragraph()
        p.style = self._pstyle(doc, ns, 'Report_H2')
        p.add_run(text)

    def _add_h3(self, doc, ns, text):
        """添加三级标题"""
        p = doc.add_paragraph()
        p.style = self._pstyle(doc, ns, 'Report_H3')
        p.add_run(text)

    def _add_body(self, doc, ns, text):
        """添加正文（带首行缩进）"""
        if not text.strip():
            return
        p = doc.add_paragraph()
        p.style = self._pstyle(doc, ns, 'Report_Body')
        p.add_run(text.strip())

    def _add_field_info(self, doc, ns, name, value):
        """添加试验基础信息（标签加粗、紧凑无缩进）"""
        if not value:
            return
        p = doc.add_paragraph()
        p.style = self._pstyle(doc, ns, 'Report_Field_Info')
        run = p.add_run(f'{name}：')
        run.font.bold = True
        run.font.color.rgb = self._REPORT_STYLES['field_label_color']
        p.add_run(str(value))

    def _add_step_item(self, doc, ns, index, title, desc):
        """添加有序试验步骤"""
        p = doc.add_paragraph()
        p.style = self._pstyle(doc, ns, 'Report_Body')
        p.paragraph_format.first_line_indent = None  # 取消缩进
        run = p.add_run(f'{index}. ')
        run.font.bold = True
        if title:
            run2 = p.add_run(f'{title}：')
            run2.font.bold = True
        if desc:
            p.add_run(desc)

    def _add_image_caption(self, doc, ns, fig_num, caption_text):
        """添加图片题注（图X 名称，居中#666666小五）"""
        p = doc.add_paragraph()
        p.style = self._pstyle(doc, ns, 'Report_Image_Caption')
        p.add_run(f'图{fig_num}  {caption_text}')

    def _add_endnote(self, doc, ns, text):
        """添加尾注"""
        p = doc.add_paragraph()
        p.style = self._pstyle(doc, ns, 'Report_Endnote')
        p.add_run(text)

    def _add_record_section(self, doc, ns, record, record_index, img_quality, fig_counter):
        """完整单条记录布局"""
        rec_title = record.get('title', f'记录{record_index + 1}')
        self._add_h2(doc, ns, f'第{record_index + 1}条记录：{rec_title}')

        # 试验基础信息区块
        self._add_field_info(doc, ns, '试验日期', record.get('date', ''))
        self._add_field_info(doc, ns, '试验人员', record.get('personnel', ''))
        self._add_field_info(doc, ns, '试验地点', record.get('location', ''))
        self._add_field_info(doc, ns, '环境条件', record.get('conditions', ''))
        self._add_field_info(doc, ns, '模板类型', record.get('template', ''))

        # 试验内容
        content = record.get('content', '').strip()
        if content:
            self._add_h3(doc, ns, '试验内容')
            for line in content.split('\n'):
                self._add_body(doc, ns, line)

        # 标准化有序试验步骤（每一步包含自己的数据与截图）
        steps = record.get('steps', [])
        if steps:
            self._add_h3(doc, ns, '试验步骤')
            img_width = self._IMG_WIDTH_MAP.get(img_quality, Inches(5.0))
            for i, step in enumerate(steps):
                st = step.get('title', '').strip()
                sd = step.get('desc', '').strip()
                step_data = step.get('data', '').strip()
                step_images = step.get('images', [])
                if st or sd:
                    self._add_step_item(doc, ns, i + 1, st, sd)
                elif step_data or step_images:
                    self._add_step_item(doc, ns, i + 1, '未命名步骤', '')

                if step_data:
                    self._add_h3(doc, ns, f'步骤{i + 1}数据记录')
                    for line in step_data.split('\n'):
                        if line.strip():
                            self._add_body(doc, ns, line)

                if step_images:
                    self._add_h3(doc, ns, f'步骤{i + 1}截图附件')
                    for img_item in step_images:
                        if isinstance(img_item, str):
                            img_path = img_item
                            img_caption = ''
                        elif isinstance(img_item, dict):
                            img_path = img_item.get('path', '')
                            img_caption = img_item.get('caption', '')
                        else:
                            continue
                        if not (os.path.exists(img_path) and os.path.isfile(img_path)):
                            continue
                        try:
                            doc.add_picture(img_path, width=img_width)
                            last_p = doc.paragraphs[-1]
                            last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            fig_counter[0] += 1
                            caption_text = img_caption if img_caption else os.path.basename(img_path)
                            self._add_image_caption(doc, ns, fig_counter[0], caption_text)
                        except Exception as e:
                            logger.warning(f'图片[{img_path}]插入失败: {e}')

        # 兼容旧版记录：旧数据可能仍保存在记录级 data/images
        legacy_data = record.get('data', '').strip()
        if legacy_data:
            self._add_h3(doc, ns, '未归属步骤数据记录')
            for line in legacy_data.split('\n'):
                if line.strip():
                    self._add_body(doc, ns, line)

        legacy_images = record.get('images', [])
        if legacy_images:
            self._add_h3(doc, ns, '未归属步骤截图附件')
            img_width = self._IMG_WIDTH_MAP.get(img_quality, Inches(5.0))
            for img_item in legacy_images:
                if isinstance(img_item, str):
                    img_path = img_item
                    img_caption = ''
                elif isinstance(img_item, dict):
                    img_path = img_item.get('path', '')
                    img_caption = img_item.get('caption', '')
                else:
                    continue
                if not (os.path.exists(img_path) and os.path.isfile(img_path)):
                    continue
                try:
                    doc.add_picture(img_path, width=img_width)
                    last_p = doc.paragraphs[-1]
                    last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    fig_counter[0] += 1
                    caption_text = img_caption if img_caption else os.path.basename(img_path)
                    self._add_image_caption(doc, ns, fig_counter[0], caption_text)
                except Exception as e:
                    logger.warning(f'图片[{img_path}]插入失败: {e}')

        # 试验结论与问题汇总
        conclusion = record.get('conclusion', '').strip()
        if conclusion:
            self._add_h3(doc, ns, '结论与问题')
            for line in conclusion.split('\n'):
                if line.strip():
                    self._add_body(doc, ns, line)

    # ---- 导出辅助 ----

    def _show_tree_context_menu(self, pos):
        """显示树右键菜单"""
        item = self.tree_projects.itemAt(pos)
        if item:
            self.tree_projects.setCurrentItem(item)
            self.tree_context_menu.exec_(self.tree_projects.mapToGlobal(pos))

    def _show_export_menu(self):
        """显示导出下拉菜单"""
        self.export_menu.exec_(self.btn_export_word.mapToGlobal(
            self.btn_export_word.rect().bottomLeft()))

    def _export_tree_selected(self, level):
        """从树右键菜单触发导出"""
        current = self.tree_projects.currentItem()
        if not current:
            QMessageBox.information(self, '提示', '请在左侧选择要导出的项目/任务/记录')
            return
        data = current.data(Qt.UserRole)
        if not data:
            return

        if level == 'record' and data[0] == 'record':
            _, pid, tid, rid = data
            self._load_record(pid, tid, rid)
            self._export_to_word('record', pid, tid, rid)
        elif level == 'task' and data[0] == 'task':
            _, pid, tid = data
            self._export_to_word('task', pid, tid)
        elif level == 'project' and data[0] == 'project':
            pid = data[1]
            self._export_to_word('project', pid)
        elif data[0] in ('record', 'task', 'project') and data[0] != level:
            QMessageBox.information(self, '提示', f'请选择对应的层级：\n当前选中的是{data[0]}，但您点击了{level}')
        else:
            QMessageBox.information(self, '提示', '请先在左侧选择对应的项目/任务/记录')

    def _export_to_word(self, level='record', pid=None, tid=None, rid=None):
        """多级导出Word文档 - 重构版（全部调用命名样式，无零散硬编码）"""
        if Document is None:
            QMessageBox.warning(
                self,
                '缺少依赖',
                '导出 Word 需要安装 python-docx。\n请执行：pip install python-docx'
            )
            return

        # ---- 导出设置对话框 ----
        dlg = QDialog(self)
        dlg.setWindowTitle('导出设置')
        dlg.setMinimumWidth(420)
        dlg_layout = QFormLayout(dlg)

        combo_style = QComboBox()
        combo_style.addItems(['📄 标准正式（宋体/黑体/标准格式）', '🌿 简洁清爽（现代/留白）', '📦 紧凑密集（省纸/高密度）'])
        dlg_layout.addRow('报告风格:', combo_style)

        chk_cover = QCheckBox()
        chk_cover.setChecked(True)
        dlg_layout.addRow('包含封面:', chk_cover)

        chk_page_num = QCheckBox()
        chk_page_num.setChecked(True)
        dlg_layout.addRow('包含页码:', chk_page_num)

        chk_toc = QCheckBox()
        chk_toc.setChecked(False)
        dlg_layout.addRow('自动目录:', chk_toc)

        combo_img = QComboBox()
        combo_img.addItems(['高清', '标准', '压缩'])
        dlg_layout.addRow('图片质量:', combo_img)

        edit_watermark = QLineEdit()
        edit_watermark.setPlaceholderText('不加水印留空')
        dlg_layout.addRow('水印文字:', edit_watermark)

        btn_box = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        btn_box.accepted.connect(dlg.accept)
        btn_box.rejected.connect(dlg.reject)
        dlg_layout.addRow(btn_box)

        if dlg.exec_() != QDialog.Accepted:
            return

        style_name = combo_style.currentText()
        include_cover = chk_cover.isChecked()
        include_page_num = chk_page_num.isChecked()
        include_toc = chk_toc.isChecked()
        img_quality = combo_img.currentText()
        watermark_text = edit_watermark.text().strip()

        # 确定导出范围
        if level == 'record':
            if pid is None or tid is None or rid is None:
                if not all([self.current_project_id, self.current_task_id, self.current_record_id]):
                    QMessageBox.information(self, '提示', '没有正在编辑的记录')
                    return
                self._save_current_record()
                pid = self.current_project_id
                tid = self.current_task_id
                rid = self.current_record_id
            try:
                recs_to_export = {rid: self.projects[pid]['tasks'][tid]['records'][rid]}
            except KeyError:
                QMessageBox.warning(self, '错误', '无法获取记录数据')
                return

        elif level == 'task':
            if pid is None or tid is None:
                current = self.tree_projects.currentItem()
                if current:
                    data = current.data(Qt.UserRole)
                    if data and data[0] == 'task':
                        _, pid, tid = data
            if not pid or not tid:
                QMessageBox.information(self, '提示', '请先在左侧选中一个任务')
                return
            try:
                recs_to_export = self.projects[pid]['tasks'][tid].get('records', {})
            except KeyError:
                QMessageBox.warning(self, '错误', '无法获取任务数据')
                return
            if not recs_to_export:
                QMessageBox.information(self, '提示', '该任务下没有记录可导出')
                return

        elif level == 'project':
            if pid is None:
                current = self.tree_projects.currentItem()
                if current:
                    data = current.data(Qt.UserRole)
                    if data and data[0] == 'project':
                        pid = data[1]
            if not pid:
                QMessageBox.information(self, '提示', '请先在左侧选中一个项目')
                return
            recs_to_export = {}
            try:
                for t_id in self.projects[pid].get('tasks', {}):
                    for r_id, rec in self.projects[pid]['tasks'][t_id].get('records', {}).items():
                        recs_to_export[(t_id, r_id)] = rec
            except KeyError:
                pass
            if not recs_to_export:
                QMessageBox.information(self, '提示', '该项目下没有记录可导出')
                return

        # 默认文件名
        if level == 'record':
            first_rec = list(recs_to_export.values())[0]
            default_name = first_rec.get('title', '试验记录').replace('/', '_').replace('\\', '_')
            default_name = f'{default_name}.docx'
        elif level == 'task':
            task_name = self.projects[pid]['tasks'][tid].get('name', '任务')
            default_name = f'{task_name}_汇总.docx'
        else:
            proj_name = self.projects[pid].get('name', '项目')
            default_name = f'{proj_name}_完整报告.docx'

        filepath, _ = QFileDialog.getSaveFileName(
            self, '导出Word文档', os.path.join(DATA_DIR, default_name), 'Word文档 (*.docx)')
        if not filepath:
            return

        try:
            # ---- 初始化文档 ----
            doc = Document()

            # ---- 页面统一设置 ----
            sec = doc.sections[0]
            sec.top_margin = self._REPORT_STYLES['page_top_margin']
            sec.bottom_margin = self._REPORT_STYLES['page_bottom_margin']
            sec.left_margin = self._REPORT_STYLES['page_left_margin']
            sec.right_margin = self._REPORT_STYLES['page_right_margin']

            # ---- 创建内置命名样式 ----
            ns = self._create_named_styles(doc)

            # ---- 封面标题文字 ----
            if level == 'project':
                proj_name = self.projects[pid].get('name', '项目')
                title_text = f'{proj_name} - 完整试验报告'
            elif level == 'task':
                task_name = self.projects[pid]['tasks'][tid].get('name', '任务')
                title_text = f'{task_name} - 汇总报告'
            else:
                first_rec = list(recs_to_export.values())[0]
                title_text = first_rec.get('title', '试验记录')

            # ---- 封面 ----
            self._add_cover_page(doc, ns, title_text, include_cover)

            # ---- 目录（可选） ----
            if include_toc:
                p = doc.add_paragraph()
                p.style = doc.styles[ns['CoverTitle']]
                run = p.add_run('目录')
                run.font.size = Pt(14)
                # Word 自动目录字段
                fldChar1 = OxmlElement('w:fldChar')
                fldChar1.set(qn('w:fldCharType'), 'begin')
                instrText = OxmlElement('w:instrText')
                instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
                fldChar2 = OxmlElement('w:fldChar')
                fldChar2.set(qn('w:fldCharType'), 'separate')
                fldChar3 = OxmlElement('w:fldChar')
                fldChar3.set(qn('w:fldCharType'), 'end')
                run2 = p.add_run()
                run2._element.append(fldChar1)
                run2._element.append(instrText)
                run2._element.append(fldChar2)
                run2._element.append(fldChar3)
                doc.add_page_break()

            # ---- 写入所有记录 ----
            rec_idx = 0
            fig_counter = [0]  # 全局图号计数器（跨记录连续编号）
            for key, rec in recs_to_export.items():
                if rec_idx > 0:
                    doc.add_page_break()
                self._add_record_section(doc, ns, rec, rec_idx, img_quality, fig_counter)
                rec_idx += 1

            # ---- 生成页眉（项目/任务名称，宋体小五中度灰居中） ----
            header_text = ''
            if level == 'project':
                header_text = self.projects[pid].get('name', '')
            elif level == 'task':
                header_text = self.projects[pid]['tasks'][tid].get('name', '')
            else:
                first_rec = list(recs_to_export.values())[0]
                header_text = first_rec.get('title', '')
            if header_text:
                header = sec.header
                hp = header.paragraphs[0]
                hp.text = header_text
                hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in hp.runs:
                    run.font.size = self._REPORT_STYLES['header_size']
                    run.font.color.rgb = self._REPORT_STYLES['header_color']
                    self._set_cn_font(run, self._REPORT_STYLES['header_font'])

            # ---- 页码（页脚，第 X 页 共 Y 页） ----
            if include_page_num:
                footer = sec.footer
                fp = footer.paragraphs[0]
                fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = fp.add_run('第 ')
                run.font.size = self._REPORT_STYLES['footer_size']
                run.font.color.rgb = self._REPORT_STYLES['footer_color']
                self._set_cn_font(run, self._REPORT_STYLES['footer_font'])

                # PAGE 域
                run2 = fp.add_run()
                fldChar1 = OxmlElement('w:fldChar')
                fldChar1.set(qn('w:fldCharType'), 'begin')
                instrText = OxmlElement('w:instrText')
                instrText.text = ' PAGE '
                fldChar2 = OxmlElement('w:fldChar')
                fldChar2.set(qn('w:fldCharType'), 'separate')
                fldChar3 = OxmlElement('w:fldChar')
                fldChar3.set(qn('w:fldCharType'), 'end')
                run2._element.append(fldChar1)
                run2._element.append(instrText)
                run2._element.append(fldChar2)
                run2._element.append(fldChar3)

                run3 = fp.add_run(' 页 共 ')
                run3.font.size = self._REPORT_STYLES['footer_size']
                run3.font.color.rgb = self._REPORT_STYLES['footer_color']
                self._set_cn_font(run3, self._REPORT_STYLES['footer_font'])

                run4 = fp.add_run()
                fldChar1b = OxmlElement('w:fldChar')
                fldChar1b.set(qn('w:fldCharType'), 'begin')
                instrText2 = OxmlElement('w:instrText')
                instrText2.text = ' NUMPAGES '
                fldChar2b = OxmlElement('w:fldChar')
                fldChar2b.set(qn('w:fldCharType'), 'separate')
                fldChar3b = OxmlElement('w:fldChar')
                fldChar3b.set(qn('w:fldCharType'), 'end')
                run4._element.append(fldChar1b)
                run4._element.append(instrText2)
                run4._element.append(fldChar2b)
                run4._element.append(fldChar3b)

            # ---- 水印（可选，浅灰色底层） ----
            if watermark_text:
                for p in doc.paragraphs:
                    for run in p.runs:
                        if run.text.strip():
                            run.font.color.rgb = self._REPORT_STYLES['watermark_color']

            # ---- 报告尾注 ----
            doc.add_paragraph()
            total_records = len(recs_to_export)
            now_str = datetime.datetime.now().strftime('%Y-%m-%d %H:%M')
            self._add_endnote(doc, ns, f'生成工具：试验记录助手 | 总记录数：{total_records} 条 | 导出时间：{now_str}')
            self._add_endnote(doc, ns, '— 本报告由「试验记录助手」自动生成 —')

            # ---- 保存 ----
            doc.save(filepath)
            logger.info(f'Word导出成功 [{style_name}]: {filepath}')
            QMessageBox.information(self, '导出成功', f'Word文档已保存到:\n{filepath}')

        except Exception as e:
            logger.error(f'Word导出失败: {e}')
            QMessageBox.critical(self, '导出失败', f'导出Word文档时发生错误:\n{str(e)}')

    # ==================== 数据持久化 ====================

    def _save_data(self):
        """保存数据到本地JSON文件"""
        filepath = os.path.join(DATA_DIR, 'experiments.json')
        try:
            with open(filepath, 'w', encoding='utf-8') as f:
                json.dump(self.projects, f, ensure_ascii=False, indent=2)
        except Exception as e:
            logger.error(f'保存数据失败: {e}')

    def _load_data(self):
        """从本地JSON文件加载数据"""
        filepath = os.path.join(DATA_DIR, 'experiments.json')
        if os.path.exists(filepath):
            try:
                with open(filepath, 'r', encoding='utf-8') as f:
                    self.projects = json.load(f)
                logger.info(f'已加载试验数据 ({len(self.projects)} 个项目)')
            except Exception as e:
                logger.error(f'加载数据失败: {e}')
                self.projects = {}
        else:
            self.projects = {}

    def statusBar(self):
        parent = self.parent()
        while parent is not None:
            if hasattr(parent, 'statusBar'):
                return parent.statusBar()
            parent = parent.parent()
        return None
